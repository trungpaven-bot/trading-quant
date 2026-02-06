
import os
from docx import Document
import re

def read_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
                        
        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading {file_path}: {e}"

def extract_info():
    # The G drive path provided by user
    input_dir = r"G:\My Drive\60_PROJECT_ROOT (AI)\00_INPUT"
    
    if not os.path.exists(input_dir):
        print(f"Cannot access {input_dir}")
        return

    files = [f for f in os.listdir(input_dir) if f.endswith('.docx')]
    
    print("--- START EXTRACTION ---")
    
    for file in files:
        print(f"\n>>> FILE: {file}")
        path = os.path.join(input_dir, file)
        text = read_docx(path)
        
        # Heuristic keywords
        targets = [
            "Tên công ty", "Công ty TNHH", "Mã số thuế", "MST", 
            "Địa chỉ", "Đại diện", "Chức vụ", "Vốn điều lệ", "Vốn đầu tư", 
            "Tên dự án", "Địa điểm thực hiện", "Diện tích"
        ]
        
        lines = text.split('\n')
        for i, line in enumerate(lines[:500]): # Scan first 500 lines/paragraphs
            for t in targets:
                if t.lower() in line.lower():
                    # Print the finding and maybe next line if it looks like a value
                    print(f"HIT [{t}]: {line.strip()}")
                    if i + 1 < len(lines):
                         print(f"    Possible Value: {lines[i+1].strip()}")

if __name__ == "__main__":
    extract_info()
