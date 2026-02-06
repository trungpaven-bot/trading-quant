
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import datetime
import os

# --- TARGET OUTPUT ---
OUTPUT_DIR = r"G:\My Drive\60_PROJECT_ROOT (AI)\03_OUTPUT"

# --- CONFIGURATION FROM MEMORY ---
PROJECT_NAME = "Kho LNG Nam Tân Tập"
PROJECT_LOCATION = "KCN Nam Tân Tập, xã Tân Tập, huyện Cần Giuộc, tỉnh Long An"
COMPANY_NAME = "CÔNG TY TNHH SAIGONTEL LONG AN"
REP_NAME = "Bà NGUYỄN CẨM PHƯƠNG"
REP_ROLE = "Tổng Giám đốc"
ADDRESS = "Ấp Tân Đông, xã Tân Tập, huyện Cần Giuộc, tỉnh Long An"
TAX_ID = "1101977241"

# --- HELPER FUNCTIONS ---
def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.first_line_indent = Cm(0.5)

def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, header=False):
    p = doc.add_paragraph()
    p.alignment = align
    if header:
        p.paragraph_format.first_line_indent = Cm(0)
    
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def create_mou_powerplant():
    doc = Document()
    set_style(doc)
    
    # Page Setup
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # 1. HEADER
    add_para(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, header=True)
    p = add_para(doc, "Độc lập – Tự do – Hạnh phúc", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, header=True)
    p.runs[0].underline = True
    add_para(doc, "----------------o0o----------------", align=WD_ALIGN_PARAGRAPH.CENTER, header=True)
    
    doc.add_paragraph()

    # 2. TITLE
    p = add_para(doc, "THỎA THUẬN NGUYÊN TẮC TIÊU THỤ KHÍ (GSA)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, header=True)
    p.runs[0].font.size = Pt(16)
    
    p = add_para(doc, "(HEADS OF AGREEMENT)", bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, header=True)
    
    p = add_para(doc, "V/v: Cung cấp khí tái hóa cho Nhà máy điện Long An I & II", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, header=True)
    
    doc.add_paragraph()
    
    # 3. BASES
    bases = [
        "Căn cứ Bộ luật Dân sự số 91/2015/QH13;",
        "Căn cứ Luật Thương mại số 36/2005/QH11;",
        "Căn cứ Quy hoạch điện VIII đã được phê duyệt;",
        "Căn cứ nhu cầu vận hành của Nhà máy điện Long An I & II và năng lực cung cấp của Kho LNG Nam Tân Tập."
    ]
    for b in bases:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5)
        run = p.add_run("- " + b)
        run.italic = True

    doc.add_paragraph()
    
    # 4. PARTIES
    today = datetime.date.today().strftime("ngày %d tháng %m năm %Y")
    add_para(doc, f"Hôm nay, {today}, tại Long An, chúng tôi gồm có:")
    
    # Party A (SELLER)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("BÊN A (BÊN BÁN): " + COMPANY_NAME).bold = True
    
    info_a = [
        f"Địa chỉ: {ADDRESS}",
        f"Đại diện: {REP_NAME} – Chức vụ: {REP_ROLE}",
        "Vai trò: Chủ đầu tư dự án Kho LNG Nam Tân Tập, đơn vị cung cấp nhiên liệu khí."
    ]
    for info in info_a:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5)
        p.add_run("- " + info)

    doc.add_paragraph()

    # Party B (BUYER)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("BÊN B (BÊN MUA): [TÊN CHỦ ĐẦU TƯ NMĐ LONG AN I & II]").bold = True
    
    info_b = [
        "Địa chỉ: [...]",
        "Đại diện: Ông/Bà [...] – Chức vụ: [...]",
        "Vai trò: Chủ đầu tư Nhà máy điện Long An I & II, đơn vị tiêu thụ nhiên liệu khí."
    ]
    for info in info_b:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5)
        p.add_run("- " + info)
        
    doc.add_paragraph()
    add_para(doc, "Hai bên thống nhất ký kết Bản Thỏa thuận nguyên tắc (HOA) để định hướng đàm phán Hợp đồng Mua bán khí (GSA) với các nội dung chính sau:")

    # 5. ARTICLES
    # Article 1
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("ĐIỀU 1. PHẠM VI HỢP TÁC").bold = True
    add_para(doc, f"Bên A cam kết cung cấp và Bên B cam kết tiêu thụ khí thiên nhiên hóa lỏng đã tái hóa (Regasified LNG) từ {PROJECT_NAME} để phục vụ vận hành thương mại cho Nhà máy điện Long An I & II.")
    add_para(doc, "Điểm giao nhận khí (Delivery Point): Tại trạm đo đếm khí đặt tại hàng rào Nhà máy điện Long An I & II.")

    # Article 2
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("ĐIỀU 2. SẢN LƯỢNG VÀ CÔNG SUẤT").bold = True
    add_para(doc, "2.1. Tổng công suất dự kiến NMĐ Long An I & II: 3.000 MW.")
    add_para(doc, "2.2. Nhu cầu khí dự kiến: Khoảng 1.000.000 - 1.500.000 tấn LNG/năm (sẽ được xác định chính xác theo, chế độ vận hành của Nhà máy điện).")
    add_para(doc, "2.3. Bên A cam kết ưu tiên tối đa công suất của Kho LNG Nam Tân Tập để đảm bảo nhiên liệu cho Bên B.")

    # Article 3
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("ĐIỀU 3. CƠ CHẾ GIÁ KHÍ (PRICE FORMULA)").bold = True
    add_para(doc, "Giá bán khí tại Điểm giao nhận sẽ được tính theo công thức chuyển ngang (Pass-through) cộng chi phí hợp lý:")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("P_gas = P_import + C_regas + C_transport").bold = True
    
    add_para(doc, "Trong đó:")
    add_para(doc, "- P_import: Giá LNG nhập khẩu theo thị trường quốc tế.")
    add_para(doc, "- C_regas: Phí dịch vụ kho cảng và tái hóa khí (Tariff).")
    add_para(doc, "- C_transport: Phí vận chuyển khí qua đường ống (nếu có).")

    # Article 4
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("ĐIỀU 4. TIẾN ĐỘ THỰC HIỆN").bold = True
    add_para(doc, "Hai bên cam kết phối hợp chặt chẽ để đồng bộ tiến độ:")
    add_para(doc, "- Tiến độ Kho LNG Nam Tân Tập: Dự kiến vận hành thương mại (COD) vào năm [20..].")
    add_para(doc, "- Tiến độ Nhà máy điện Long An I & II: Dự kiến vận hành thương mại (COD) vào năm [20..].")

    # Article 5
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("ĐIỀU 5. ĐIỀU KHOẢN CHUNG").bold = True
    add_para(doc, "Bản thỏa thuận này thể hiện thiện chí hợp tác và là tiền đề pháp lý để hai bên báo cáo Cơ quan chức năng (Bộ Công Thương, UBND Tỉnh Long An) về phương án đảm bảo nhiên liệu cho dự án điện.")
    
    doc.add_paragraph()
    doc.add_paragraph()

    # 6. SIGNATURES
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    
    c1 = table.cell(0, 0)
    c2 = table.cell(0, 1)
    
    # Sig A
    p = c1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ĐẠI DIỆN BÊN A (BÊN BÁN)").bold = True
    p = c1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("TỔNG GIÁM ĐỐC").bold = True
    
    # Space for sign
    c1.add_paragraph()
    c1.add_paragraph()
    c1.add_paragraph()
    
    p = c1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(REP_NAME).bold = True

    # Sig B
    p = c2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ĐẠI DIỆN BÊN B (BÊN MUA)").bold = True
    p = c2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("TỔNG GIÁM ĐỐC").bold = True
    
    c2.add_paragraph()
    c2.add_paragraph()
    c2.add_paragraph()
    
    p = c2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("[TÊN ĐẠI DIỆN BÊN B]").bold = True

    # Save
    target_dir = OUTPUT_DIR
    if not os.path.exists(target_dir):
        print(f"Warning: {target_dir} does not exist. Saving to local.")
        target_dir = "."
        
    date_str = datetime.date.today().strftime("%Y-%m-%d")
    filename = f"{date_str}_MOU_Tieu thu khi_NM Dien Long An.docx"
    output_path = os.path.join(target_dir, filename)
    
    doc.save(output_path)
    print(f"File created: {output_path}")

if __name__ == "__main__":
    create_mou_powerplant()
