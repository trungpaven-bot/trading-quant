
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import datetime
import os

# --- TARGET OUTPUT ---
OUTPUT_DIR = r"G:\My Drive\60_PROJECT_ROOT (AI)\03_OUTPUT"

# --- CONFIGURATION FROM MEMORY ---
PROJECT_NAME = "Kho LNG Nam Tân Tập"
PROJECT_LOCATION = "KCN Nam Tân Tập, xã Tân Tập, huyện Cần Giuộc, tỉnh Long An"
COMPANY_NAME = "CÔNG TY TNHH SAIGONTEL LONG AN"
REP_NAME = "Bà NGUYỄN CẨM PHƯƠNG"
REP_ROLE = "Tổng Giám đốc"
ADDRESS = "Ấp Tân Đông, xã Tân Tập, huyện Cần Giuộc, tỉnh Long An"
TAX_ID = "1101977241"

# --- HELPER FUNCTIONS ---
def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.first_line_indent = Cm(0.5)

def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, header=False):
    p = doc.add_paragraph()
    p.alignment = align
    if header:
        p.paragraph_format.first_line_indent = Cm(0)
    
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def create_mou():
    doc = Document()
    set_style(doc)
    
    # Page Setup
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # 1. HEADER
    add_para(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, header=True)
    p = add_para(doc, "Độc lập – Tự do – Hạnh phúc", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, header=True)
    p.runs[0].underline = True
    add_para(doc, "----------------o0o----------------", align=WD_ALIGN_PARAGRAPH.CENTER, header=True)
    
    doc.add_paragraph()

    # 2. TITLE
    p = add_para(doc, "THỎA THUẬN NGUYÊN TẮC CUNG CẤP KHÍ (LNG)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, header=True)
    p.runs[0].font.size = Pt(16)
    
    p = add_para(doc, "MEMORANDUM OF UNDERSTANDING (MOU)", bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, header=True)
    
    doc.add_paragraph()
    
    # 3. BASES
    bases = [
        "Căn cứ Bộ luật Dân sự số 91/2015/QH13;",
        "Căn cứ Luật Thương mại số 36/2005/QH11;",
        "Căn cứ nhu cầu và năng lực của hai Bên."
    ]
    for b in bases:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5)
        run = p.add_run("- " + b)
        run.italic = True

    doc.add_paragraph()
    
    # 4. PARTIES
    today = datetime.date.today().strftime("ngày %d tháng %m năm %Y")
    add_para(doc, f"Hôm nay, {today}, tại Long An, chúng tôi gồm có:")
    
    # Party A
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("BÊN A (BÊN MUA): " + COMPANY_NAME).bold = True
    
    info_a = [
        f"Địa chỉ: {ADDRESS}",
        f"Mã số thuế: {TAX_ID}",
        f"Đại diện: {REP_NAME} – Chức vụ: {REP_ROLE}"
    ]
    for info in info_a:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5)
        p.add_run("- " + info)

    doc.add_paragraph()

    # Party B
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("BÊN B (BÊN CUNG CẤP): [TÊN ĐỐI TÁC CUNG CẤP KHÍ]").bold = True
    
    info_b = [
        "Địa chỉ: [Điền địa chỉ đối tác]",
        "Mã số thuế/GPKD: [Điền số]",
        "Đại diện: Ông/Bà [Tên đại diện] – Chức vụ: [Chức vụ]"
    ]
    for info in info_b:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5)
        p.add_run("- " + info)
        
    doc.add_paragraph()
    add_para(doc, "Sau khi bàn bạc, hai bên thống nhất ký kết Bản Thỏa thuận nguyên tắc với các điều khoản sau:")

    # 5. ARTICLES
    # Article 1
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("ĐIỀU 1. MỤC TIÊU THỎA THUẬN").bold = True
    add_para(doc, f"Bên B đồng ý cung cấp và Bên A đồng ý tiếp nhận nguồn khí hóa lỏng (LNG) để phục vụ vận hành thương mại cho {PROJECT_NAME}.")
    add_para(doc, f"Địa điểm giao hàng: Tại {PROJECT_LOCATION}.")

    # Article 2
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("ĐIỀU 2. KHỐI LƯỢNG VÀ TIẾN ĐỘ (DỰ KIẾN)").bold = True
    add_para(doc, "2.1. Khối lượng dự kiến: Khoảng 1.000.000 tấn LNG/năm (theo công suất thiết kế của Kho).")
    add_para(doc, "2.2. Thời gian bắt đầu cung cấp: Dự kiến từ Quý [Strike]... năm 20[..] (ngay khi Kho LNG đi vào vận hành).")
    add_para(doc, "2.3. Quy cách kỹ thuật: LNG đáp ứng tiêu chuẩn quốc tế và quy định hiện hành của Việt Nam.")

    # Article 3
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("ĐIỀU 3. GIÁ CẢ VÀ PHƯƠNG THỨC THANH TOÁN").bold = True
    add_para(doc, "Giá bán khí sẽ được xác định dựa trên công thức giá thị trường thế giới tại thời điểm giao hàng cộng với các chi phí vận chuyển, bảo hiểm hợp lý. Chi tiết sẽ được quy định cụ thể trong Hợp đồng Mua bán Khí (GSPA) chính thức.")

    # Article 4
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("ĐIỀU 4. TRÁCH NHIỆM CÁC BÊN").bold = True
    add_para(doc, "4.1. Trách nhiệm Bên A: Đảm bảo tiến độ xây dựng Kho LNG để sẵn sàng tiếp nhận khí; Thanh toán đúng hạn.")
    add_para(doc, "4.2. Trách nhiệm Bên B: Đảm bảo nguồn cung ổn định, chất lượng khí và hỗ trợ kỹ thuật trong quá trình nhập hàng.")

    # Article 5
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.add_run("ĐIỀU 5. ĐIỀU KHOẢN CHUNG").bold = True
    add_para(doc, "- Thỏa thuận này là cơ sở để hai bên tiến hành đàm phán Hợp đồng Mua bán Khí (GSPA) chính thức.")
    add_para(doc, "- Thỏa thuận có hiệu lực kể từ ngày ký.")
    
    doc.add_paragraph()
    doc.add_paragraph()

    # 6. SIGNATURES
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    
    c1 = table.cell(0, 0)
    c2 = table.cell(0, 1)
    
    # Sig A
    p = c1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ĐẠI DIỆN BÊN A").bold = True
    p = c1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("TỔNG GIÁM ĐỐC").bold = True
    
    # Space for sign
    c1.add_paragraph()
    c1.add_paragraph()
    c1.add_paragraph()
    
    p = c1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(REP_NAME).bold = True

    # Sig B
    p = c2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ĐẠI DIỆN BÊN B").bold = True
    p = c2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("TỔNG GIÁM ĐỐC").bold = True
    
    c2.add_paragraph()
    c2.add_paragraph()
    c2.add_paragraph()
    
    p = c2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("[TÊN ĐẠI DIỆN BÊN B]").bold = True

    # Save
    # Save
    target_dir = OUTPUT_DIR
    if not os.path.exists(target_dir):
        print(f"Warning: {target_dir} does not exist. Saving to local.")
        target_dir = "."
        
    date_str = datetime.date.today().strftime("%Y-%m-%d")
    filename = f"{date_str}_MOU_Cung cap khi_Kho LNG Nam Tan Tap.docx"
    output_path = os.path.join(target_dir, filename)
    
    doc.save(output_path)
    print(f"File created: {output_path}")

if __name__ == "__main__":
    create_mou()
