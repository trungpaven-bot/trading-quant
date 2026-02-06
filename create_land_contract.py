
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_contract():
    doc = Document()
    
    # Style settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # 1. HEADER
    p = doc.add_paragraph('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    p = doc.add_paragraph('Độc lập – Tự do – Hạnh phúc')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True
    
    doc.add_paragraph() # Spacer
    
    title = doc.add_paragraph('HỢP ĐỒNG NGUYÊN TẮC\nCHUYỂN NHƯỢNG QUYỀN SỬ DỤNG ĐẤT')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    
    intro = doc.add_paragraph('Hôm nay, ngày .... tháng .... năm 2026, tại ..................., chúng tôi gồm có:')
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # 2. PARTIES
    # Party A
    doc.add_paragraph('BÊN CHUYỂN NHƯỢNG (BÊN A): [TÊN TỔ CHỨC/CÁ NHÂN SỞ HỮU ĐẤT]', style='List Bullet')
    doc.add_paragraph('Địa chỉ trụ sở: ............................................................................................')
    doc.add_paragraph('Mã số thuế/CCCD: ........................................................................................')
    doc.add_paragraph('Đại diện bởi: Ông/Bà .................................... Chức vụ: ....................................')
    
    doc.add_paragraph()
    
    # Party B
    doc.add_paragraph('BÊN NHẬN CHUYỂN NHƯỢNG (BÊN B): [TÊN ĐẠI DIỆN LIÊN DANH / CÔNG TY DỰ ÁN DỰ KIẾN]', style='List Bullet')
    doc.add_paragraph('Đại diện cho Liên danh đầu tư Dự án Kho LNG Nam Tân Tập.')
    doc.add_paragraph('Địa chỉ: .........................................................................................................')
    doc.add_paragraph('Đại diện bởi: Ông/Bà .................................... Chức vụ: ....................................')
    
    doc.add_paragraph()
    doc.add_paragraph('Hai bên cùng thống nhất ký kết Hợp đồng nguyên tắc về việc chuyển nhượng quyền sử dụng đất để thực hiện Dự án "Kho LNG Nam Tân Tập" với các điều khoản như sau:')
    
    # 3. CONTENT
    # Article 1
    doc.add_paragraph('ĐIỀU 1. ĐỐI TƯỢNG HỢP ĐỒNG', style='List Number')
    doc.add_paragraph('1.1. Bên A đồng ý chuyển nhượng và Bên B đồng ý nhận chuyển nhượng toàn bộ quyền sử dụng đất tại thửa đất số ...., tờ bản đồ số ...., tại xã Tân Tập, huyện Cần Giuộc, tỉnh Long An.')
    doc.add_paragraph('1.2. Diện tích đất dự kiến chuyển nhượng: .................... m2 (Bằng chữ: .................................................).')
    doc.add_paragraph('1.3. Mục đích sử dụng: Thực hiện Dự án đầu tư xây dựng Kho LNG Nam Tân Tập theo Chủ trương đầu tư được cấp có thẩm quyền phê duyệt.')
    
    # Article 2
    doc.add_paragraph('ĐIỀU 2. GIÁ CHUYỂN NHƯỢNG VÀ PHƯƠNG THỨC THANH TOÁN', style='List Number')
    doc.add_paragraph('2.1. Giá chuyển nhượng (tạm tính): .................... VNĐ/m2.')
    doc.add_paragraph('Tổng giá trị chuyển nhượng dự kiến: .................... VNĐ.')
    doc.add_paragraph('2.2. Giá trên [đã/chưa] bao gồm các loại thuế, phí liên quan đến việc chuyển nhượng quyền sử dụng đất theo quy định của pháp luật.')
    doc.add_paragraph('2.3. Phương thức thanh toán: Chuyển khoản vào tài khoản chỉ định của Bên A.')
    
    # Article 3
    doc.add_paragraph('ĐIỀU 3. TIẾN ĐỘ THỰC HIỆN VÀ ĐIỀU KIỆN CHUYỂN NHƯỢNG', style='List Number')
    doc.add_paragraph('3.1. Đặt cọc: Bên B sẽ đặt cọc cho Bên A số tiền là .................... VNĐ trong vòng .... ngày kể từ ngày ký Hợp đồng này để đảm bảo nghĩa vụ.')
    doc.add_paragraph('3.2. Ký Hợp đồng Chuyển nhượng chính thức (Công chứng): Hai bên cam kết tiến hành ký kết Hợp đồng chuyển nhượng quyền sử dụng đất chính thức tại Phòng Công chứng Nhà nước trong vòng .... ngày kể từ ngày Bên B (hoặc Công ty Dự án do Bên B thành lập) nhận được Giấy chứng nhận đăng ký đầu tư/Quyết định chấp thuận chủ trương đầu tư cho Dự án.')
    doc.add_paragraph('3.3. Trong trường hợp Bên B thành lập Pháp nhân mới (Công ty Dự án) để thực hiện dự án, Bên A đồng ý ký Hợp đồng chuyển nhượng chính thức trực tiếp với Pháp nhân mới này. Mọi quyền lợi và nghĩa vụ của Bên B trong Hợp đồng nguyên tắc này sẽ được chuyển giao kế thừa cho Pháp nhân mới.')
    
    # Article 4
    doc.add_paragraph('ĐIỀU 4. TRÁCH NHIỆM CỦA CÁC BÊN', style='List Number')
    doc.add_paragraph('4.1. Trách nhiệm của Bên A:')
    doc.add_paragraph('- Đảm bảo đất thuộc quyền sử dụng hợp pháp, không có tranh chấp, không bị kê biên, không bị cầm cố, thế chấp tại thời điểm ký Hợp đồng chính thức.')
    doc.add_paragraph('- Hỗ trợ Bên B trong quá trình thực hiện các thủ tục pháp lý để xin cấp phép đầu tư dự án liên quan đến khu đất.')
    doc.add_paragraph('4.2. Trách nhiệm của Bên B:')
    doc.add_paragraph('- Thanh toán đầy đủ và đúng hạn theo thỏa thuận.')
    doc.add_paragraph('- Chịu trách nhiệm thực hiện các thủ tục pháp lý để triển khai Dự án theo quy định.')
    
    # Article 5
    doc.add_paragraph('ĐIỀU 5. ĐIỀU KHOẢN CHUNG', style='List Number')
    doc.add_paragraph('5.1. Hai bên cam kết thực hiện đúng các điều khoản đã thỏa thuận.')
    doc.add_paragraph('5.2. Hợp đồng nguyên tắc này có hiệu lực kể từ ngày ký và sẽ hết hiệu lực khi hai bên ký Hợp đồng chuyển nhượng chính thức hoặc có văn bản thỏa thuận hủy bỏ.')
    doc.add_paragraph('5.3. Hợp đồng này được lập thành 02 bản, mỗi bên giữ 01 bản có giá trị pháp lý như nhau.')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True 
    row_cells = table.rows[0].cells
    
    # A
    p_a = row_cells[0].add_paragraph('ĐẠI DIỆN BÊN A')
    p_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_a.runs[0].bold = True
    
    # B
    p_b = row_cells[1].add_paragraph('ĐẠI DIỆN BÊN B')
    p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_b.runs[0].bold = True
    
    # Save
    output_path = r'g:\My Drive\30_RESOURCES (Kho tai nguyen)\30.01_Tinh hieu qua du an\02_Tra dinh muc theo TT 09 va 12\2026-01-30_HDNT_ChuyenNhuongDat_KhoLNG.docx'
    doc.save(output_path)
    print(f"Document saved to {output_path}")

try:
    create_contract()
except Exception as e:
    print(f"Error: {e}")
